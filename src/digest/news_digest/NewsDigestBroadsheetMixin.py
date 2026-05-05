import os
from datetime import datetime

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from utils import Log

log = Log("NewsDigestBroadsheetMixin")

# Sri Lanka flag colour palette
_CLR_MAROON  = RGBColor(0x8D, 0x15, 0x3A)  # lion-field crimson
_CLR_SAFFRON = RGBColor(0xFC, 0x8B, 0x00)  # saffron stripe
_CLR_GREEN   = RGBColor(0x00, 0x53, 0x4E)  # green stripe
_CLR_GOLD    = RGBColor(0xFF, 0xD1, 0x00)  # border / bo-leaves
_CLR_BLACK   = RGBColor(0x00, 0x00, 0x00)
_FONT        = "Ubuntu"


def _ordinal(n):
    if 11 <= n % 100 <= 13:
        return f"{n}th"
    return f"{n}{['th','st','nd','rd','th','th','th','th','th','th'][n % 10]}"


def _fmt_date(date_str):
    """Format 'YYYY-MM-DD' as 'Tuesday the 5th of March, 2026'."""
    dt = datetime.strptime(date_str, "%Y-%m-%d")
    return f"{dt.strftime('%A')} the {_ordinal(dt.day)} of {dt.strftime('%B, %Y')}"


def _r(run, pt, bold=False, color=_CLR_BLACK, italic=False):
    """Apply standard styling to a run."""
    run.font.name   = _FONT
    run.font.size   = Pt(pt)
    run.bold        = bold
    run.italic      = italic
    run.font.color.rgb = color
    return run


class NewsDigestBroadsheetMixin:
    DIR_BROADSHEETS = os.path.join("data", "broadsheets")

    _TOTAL_COLS = 4      # col 0 = Other News; cols 1-3 = main content
    _PAGE_W_CM  = 59.4   # A2 landscape width
    _PAGE_H_CM  = 42.0   # A2 landscape height
    _MARGIN_IN  = 1.0    # 1-inch padding on all sides

    _OTHER_NEWS_WORD_BUDGET = 400

    @staticmethod
    def _set_no_borders(cell):
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement("w:tcBorders")
        for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
            el = OxmlElement(f"w:{side}")
            el.set(qn("w:val"), "nil")
            tcBorders.append(el)
        tcPr.append(tcBorders)

    @staticmethod
    def _clear_table_borders(table):
        tbl  = table._tbl
        tblPr = tbl.find(qn("w:tblPr"))
        if tblPr is None:
            tblPr = OxmlElement("w:tblPr")
            tbl.insert(0, tblPr)
        tblBorders = OxmlElement("w:tblBorders")
        for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
            el = OxmlElement(f"w:{side}")
            el.set(qn("w:val"), "nil")
            tblBorders.append(el)
        tblPr.append(tblBorders)

    def _add_article_to_cell(
        self, cell, article,
        title_pt, body_pt,
        title_color=_CLR_BLACK, title_bold=True
    ):
        title_p = cell.add_paragraph()
        title_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        title_p.paragraph_format.space_before = Pt(0)
        title_p.paragraph_format.space_after  = Pt(4)
        _r(title_p.add_run(article["title"]), title_pt,
           bold=title_bold, color=title_color)

        body = article.get("body", "")
        if body:
            body_p = cell.add_paragraph()
            body_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            body_p.paragraph_format.space_before = Pt(0)
            body_p.paragraph_format.space_after  = Pt(10)
            _r(body_p.add_run(body), body_pt, color=_CLR_BLACK)

    @staticmethod
    def _split_text(text, n):
        words = text.split()
        if not words:
            return [""] * n
        chunk_size = max(1, len(words) // n)
        chunks = []
        for i in range(n):
            start = i * chunk_size
            end   = start + chunk_size if i < n - 1 else len(words)
            chunks.append(" ".join(words[start:end]))
        return chunks

    def _cap_level2_by_budget(self, articles):
        result, used = [], 0
        for article in articles:
            words = (
                len(article["title"].split())
                + len(article.get("body", "").split())
            )
            if used + words > self._OTHER_NEWS_WORD_BUDGET:
                break
            result.append(article)
            used += words
        return result

    def build_broadsheet(self, digest_article_list, ts, used_articles=None):
        doc = Document()

        # Set Ubuntu as the document default font
        doc.styles["Normal"].font.name = _FONT
        doc.styles["Normal"].paragraph_format.space_after = Pt(0)

        # A2 landscape, 1-inch margins
        section = doc.sections[0]
        section.page_width    = Cm(self._PAGE_W_CM)
        section.page_height   = Cm(self._PAGE_H_CM)
        section.left_margin   = Inches(self._MARGIN_IN)
        section.right_margin  = Inches(self._MARGIN_IN)
        section.top_margin    = Inches(self._MARGIN_IN)
        section.bottom_margin = Inches(self._MARGIN_IN)

        # --- Masthead ---
        masthead = doc.add_paragraph()
        masthead.alignment = WD_ALIGN_PARAGRAPH.CENTER
        masthead.paragraph_format.space_before = Pt(0)
        masthead.paragraph_format.space_after  = Pt(6)
        _r(masthead.add_run("SRI LANKA THIS WEEK"), 72,
           bold=True, color=_CLR_MAROON)

        # --- Date / subtitle ---
        date_p = doc.add_paragraph()
        date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_p.paragraph_format.space_before = Pt(0)
        date_p.paragraph_format.space_after  = Pt(14)

        if used_articles:
            date_strs  = sorted(a.date_str for a in used_articles)
            date_line  = (
                f"{_fmt_date(date_strs[0])}  to  {_fmt_date(date_strs[-1])}"
                f"   \u00b7   Based on {len(used_articles):,} articles"
            )
        else:
            date_line = ts

        _r(date_p.add_run(date_line), 14, italic=True, color=_CLR_SAFFRON)

        level0      = [a for a in digest_article_list if a.get("level") == 0]
        level1      = [a for a in digest_article_list if a.get("level") == 1]
        level2_all  = [a for a in digest_article_list if a.get("level") == 2]
        level2      = self._cap_level2_by_budget(level2_all)
        level2_overflow = level2_all[len(level2):]

        # --- Layout planning: overflow level-2 placement ---
        n_l1_rows              = (len(level1) + 2) // 3 if level1 else 0
        last_row_filled        = len(level1) % 3           # 0 → last row full (or no rows)
        overflow_slots_in_last = (3 - last_row_filled) % 3 if n_l1_rows else 0
        overflow_in_last_row   = level2_overflow[:overflow_slots_in_last]
        overflow_extra         = level2_overflow[overflow_slots_in_last:]
        n_extra_rows           = (len(overflow_extra) + 2) // 3 if overflow_extra else 0

        n_rows = 2 + n_l1_rows + n_extra_rows

        # --- Body font size: scale up when content is sparse ---
        total_content_rows = n_l1_rows + n_extra_rows
        if total_content_rows <= 1:
            body_pt = 16
        elif total_content_rows == 2:
            body_pt = 14
        elif total_content_rows == 3:
            body_pt = 13
        else:
            body_pt = 12

        table = doc.add_table(rows=n_rows, cols=self._TOTAL_COLS)
        self._clear_table_borders(table)

        # --- Col 0: Other News sidebar ---
        other_cell = table.cell(0, 0).merge(table.cell(n_rows - 1, 0))
        self._set_no_borders(other_cell)

        header_p = other_cell.paragraphs[0]
        header_p.paragraph_format.space_before = Pt(0)
        header_p.paragraph_format.space_after  = Pt(8)
        _r(header_p.add_run("OTHER NEWS"), 14,
           bold=True, color=_CLR_SAFFRON)

        for article in level2:
            self._add_article_to_cell(
                other_cell, article,
                title_pt=12, body_pt=body_pt,
                title_color=_CLR_SAFFRON
            )

        # --- Row 0, cols 1-3: headline title ---
        headline_cell = table.cell(0, 1).merge(table.cell(0, 3))
        self._set_no_borders(headline_cell)
        if level0:
            title_p = headline_cell.add_paragraph()
            title_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            title_p.paragraph_format.space_before = Pt(0)
            title_p.paragraph_format.space_after  = Pt(10)
            _r(title_p.add_run(level0[0]["title"]), 48,
               bold=True, color=_CLR_MAROON)

        # --- Row 1, cols 1-3: headline body in 3 columns ---
        body        = level0[0].get("body", "") if level0 else ""
        body_chunks = self._split_text(body, 3)
        for col_offset, chunk in enumerate(body_chunks):
            body_cell = table.cell(1, 1 + col_offset)
            self._set_no_borders(body_cell)
            if chunk:
                body_p = body_cell.add_paragraph()
                body_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                body_p.paragraph_format.space_before = Pt(0)
                body_p.paragraph_format.space_after  = Pt(14)
                _r(body_p.add_run(chunk), body_pt, color=_CLR_BLACK)

        # --- Rows 2+: level-1 articles, 3 per row ---
        for i, article in enumerate(level1):
            row     = 2 + i // 3
            col     = 1 + i % 3
            l1_cell = table.cell(row, col)
            self._set_no_borders(l1_cell)
            self._add_article_to_cell(
                l1_cell, article,
                title_pt=24, body_pt=body_pt,
                title_color=_CLR_GREEN
            )

        # --- Overflow level-2: fill empty slots in last level-1 row ---
        if n_l1_rows > 0 and overflow_in_last_row:
            for i, article in enumerate(overflow_in_last_row):
                col  = 1 + last_row_filled + i
                cell = table.cell(2 + n_l1_rows - 1, col)
                self._set_no_borders(cell)
                self._add_article_to_cell(
                    cell, article,
                    title_pt=12, body_pt=body_pt,
                    title_color=_CLR_SAFFRON
                )

        # --- Overflow level-2: extra rows below level-1 ---
        for i, article in enumerate(overflow_extra):
            row  = 2 + n_l1_rows + i // 3
            col  = 1 + i % 3
            cell = table.cell(row, col)
            self._set_no_borders(cell)
            self._add_article_to_cell(
                cell, article,
                title_pt=12, body_pt=body_pt,
                title_color=_CLR_SAFFRON
            )

        # --- Footer ---
        footer   = section.footer
        footer_p = footer.paragraphs[0]
        footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_p.paragraph_format.space_before = Pt(0)
        footer_p.paragraph_format.space_after  = Pt(0)
        _r(footer_p.add_run("https://github.com/nuuuwan/lk_news_digest"),
           9, italic=True, color=_CLR_MAROON)

        os.makedirs(self.DIR_BROADSHEETS, exist_ok=True)
        broadsheet_path = os.path.join(
            self.DIR_BROADSHEETS, f"broadsheet.{ts}.docx"
        )
        doc.save(broadsheet_path)
        log.info(f"Wrote broadsheet to {broadsheet_path}")
